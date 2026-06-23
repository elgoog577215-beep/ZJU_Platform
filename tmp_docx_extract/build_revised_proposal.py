from pathlib import Path
from math import atan2, cos, sin, pi

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


BASE = Path("/Users/yq/Desktop/ZJU_Platform/tmp_docx_extract/clean_base.docx")
OUT = Path("/Users/yq/Desktop/浙江大学 重点 应奇(5)-按建议修改稿.docx")
ASSET_DIR = Path("/Users/yq/Desktop/ZJU_Platform/tmp_docx_extract/assets")
ASSET_DIR.mkdir(parents=True, exist_ok=True)

FONT_CN = "/System/Library/Fonts/STHeiti Medium.ttc"
FONT_CN_LIGHT = "/System/Library/Fonts/STHeiti Light.ttc"


def font(size, bold=False):
    return ImageFont.truetype(FONT_CN if bold else FONT_CN_LIGHT, size)


def wrap(draw, text, fnt, max_width):
    lines, line = [], ""
    for ch in text:
        trial = line + ch
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width or not line:
            line = trial
        else:
            lines.append(line)
            line = ch
    if line:
        lines.append(line)
    return lines


def draw_centered(draw, box, text, fnt, fill):
    x1, y1, x2, y2 = box
    lines = wrap(draw, text, fnt, x2 - x1 - 36)
    heights = [draw.textbbox((0, 0), line, font=fnt)[3] for line in lines]
    total = sum(heights) + (len(lines) - 1) * 8
    y = y1 + (y2 - y1 - total) / 2
    for line, h in zip(lines, heights):
        w = draw.textbbox((0, 0), line, font=fnt)[2]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=fill)
        y += h + 8


def arrow(draw, start, end, fill="#2f5597", width=5):
    draw.line([start, end], fill=fill, width=width)
    ang = atan2(end[1] - start[1], end[0] - start[0])
    size = 18
    p1 = (end[0] - size * cos(ang - pi / 6), end[1] - size * sin(ang - pi / 6))
    p2 = (end[0] - size * cos(ang + pi / 6), end[1] - size * sin(ang + pi / 6))
    draw.polygon([end, p1, p2], fill=fill)


def make_route_chart(path):
    img = Image.new("RGB", (1600, 720), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(40, True)
    box_title = font(26, True)
    box_text = font(24)
    draw.text((64, 40), "技术路线图：多源数据融合驱动的校园信息智能服务与质量评估", font=title_font, fill="#17233c")

    boxes = [
        ((70, 180, 330, 360), "数据来源层", "官网通知、公众号、社团活动、竞赛讲座、志愿服务、影像资料、用户行为"),
        ((390, 180, 650, 360), "采集治理层", "自动采集、人工上传、大模型解析、字段校验、去重合并、标签标准化"),
        ((710, 180, 970, 360), "资源索引层", "活动库、AI 社区、影像库、主体档案、资源检索索引"),
        ((1030, 180, 1290, 360), "智能服务层", "全站检索、活动推荐、组织归属、数据看板、反馈记录"),
        ((550, 500, 1050, 650), "验证输出层", "线上试运行 -> 指标统计 -> 评估报告 -> 数据规范与运营制度 -> 同类高校复制参考"),
    ]

    for i, (box, head, body) in enumerate(boxes):
        fill = "#eaf2ff" if i < 4 else "#f4f7fb"
        outline = "#2f5597"
        draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=4)
        x1, y1, x2, _ = box
        draw.text((x1 + 26, y1 + 24), head, font=box_title, fill="#1f4e79")
        draw_centered(draw, (x1 + 12, y1 + 70, x2 - 12, box[3] - 18), body, box_text, "#1f2937")

    for a, b in [((330, 270), (390, 270)), ((650, 270), (710, 270)), ((970, 270), (1030, 270))]:
        arrow(draw, a, b)
    arrow(draw, (1160, 360), (960, 500))
    arrow(draw, (640, 500), (210, 360))
    draw.text((1220, 435), "运行数据反馈修正", font=font(24, True), fill="#c2410c")
    arrow(draw, (1220, 430), (620, 430), fill="#c2410c", width=4)
    img.save(path)


def make_platform_collage(path):
    srcs = [
        ("/Users/yq/Desktop/ZJU_Platform/ui-audit-screenshots/current-laptop-final/1280x720__events.png", "活动聚合与筛选"),
        ("/Users/yq/Desktop/ZJU_Platform/ui-audit-screenshots/current-laptop-final/1280x720__articles.png", "AI 社区与共建交流"),
        ("/Users/yq/Desktop/ZJU_Platform/ui-audit-screenshots/current-laptop-final/1280x720__media.png", "影像库与资料沉淀"),
    ]
    panel_w, panel_h = 500, 282
    img = Image.new("RGB", (1620, 430), "white")
    draw = ImageDraw.Draw(img)
    draw.text((36, 28), "现有平台基础截图", font=font(36, True), fill="#17233c")
    for i, (src, label) in enumerate(srcs):
        x = 36 + i * 525
        y = 94
        shot = Image.open(src).convert("RGB")
        shot.thumbnail((panel_w, panel_h), Image.Resampling.LANCZOS)
        frame = Image.new("RGB", (panel_w, panel_h), "#f8fafc")
        frame.paste(shot, ((panel_w - shot.width) // 2, (panel_h - shot.height) // 2))
        img.paste(frame, (x, y))
        draw.rectangle((x, y, x + panel_w, y + panel_h), outline="#cbd5e1", width=3)
        draw.text((x + 8, y + panel_h + 18), label, font=font(27, True), fill="#1f4e79")
    img.save(path)


def set_run_font(run, size=10.5, bold=False, color=None, name="宋体"):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def style_paragraph(paragraph, size=10.5, bold=False, align=None, color=None, font_name="宋体"):
    if align is not None:
        paragraph.alignment = align
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(21) if not bold else None
    pf.space_after = Pt(6)
    pf.line_spacing = 1.18
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color, name=font_name)


def add_before(anchor, text="", size=10.5, bold=False, align=None, color=None, font_name="宋体"):
    p = anchor.insert_paragraph_before(text)
    style_paragraph(p, size=size, bold=bold, align=align, color=color, font_name=font_name)
    return p


def add_caption_before(anchor, text):
    return add_before(anchor, text, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, font_name="宋体")


def move_before(anchor, element):
    parent = anchor._p.getparent()
    parent.insert(parent.index(anchor._p), element)


def add_image_before(doc, anchor, image_path, width_inches, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    move_before(anchor, p._p)
    if caption:
        add_caption_before(anchor, caption)
    return p


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_i == 0:
                set_cell_shading(cell, "D9EAF7")
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.1
                for r in p.runs:
                    set_run_font(r, size=9.5, bold=(row_i == 0), name="宋体")


def add_table_before(doc, anchor, rows, widths=None):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = value
            if widths:
                cell.width = Inches(widths[j])
    style_table(table)
    move_before(anchor, table._tbl)
    return table


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def find_body_para(doc, prefix, start=0):
    for i, p in enumerate(doc.paragraphs[start:], start):
        if p.text.strip().startswith(prefix):
            return i, p
    raise ValueError(f"paragraph not found: {prefix}")


def replace_section(doc, start_prefix, end_prefix, inserter):
    paras = doc.paragraphs
    start_i, _ = find_body_para(doc, start_prefix)
    end_i, anchor = find_body_para(doc, end_prefix, start_i + 1)
    inserter(anchor)
    for p in paras[start_i:end_i]:
        remove_paragraph(p)


def update_summary_table(doc):
    summary = (
        "针对当前高校校园活动、教务通知、科创竞赛、志愿服务等信息分散在官网、公众号、社群等渠道，"
        "存在数据孤岛、检索效率低、数据价值未充分挖掘等问题，本项目紧扣 A10-2 方向，开展基于多源数据融合的校园信息智能检索与教育质量评估研究。"
        "项目拟在现有平台运行基础上，完善数据采集、字段标准化、资源索引、全站检索、个性化推荐和主体归属机制，实现校园信息自动检索、挖掘、整合与精准推送；"
        "同时围绕活动参与度、资源覆盖度、检索反馈、组织活跃度等指标，构建轻量化教育质量观察与服务优化模型。"
        "成果将形成可上线验证的平台升级版本、研究报告、数据规范和运营制度，为智慧校园建设和同类高校信息治理提供可复制的实践案例。"
    )
    cell = doc.tables[0].rows[13].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run("申请课题简介（不超过500字）：")
    set_run_font(r, size=10.5, bold=True)
    p.paragraph_format.space_after = Pt(4)
    p2 = cell.add_paragraph(summary)
    style_paragraph(p2, size=10.5)


def section_three(doc, anchor):
    add_before(anchor, "三、课题研究的目的和意义", size=14, bold=True, font_name="黑体")
    add_before(anchor, (
        "本课题面向高校信息化与教育数据治理场景，对应指南 A10-2“基于多源数据融合的高校教育质量评估体系及制度创新研究”方向。"
        "当前高校校园信息长期分布在学院官网、公众号、社团通知群、赛事平台和线下公告等渠道，信息格式不统一、更新节奏不一致、历史数据难复用，学生端表现为“找不到、找不全、来不及”，管理端则难以依据真实运行数据观察校园服务质量和学生发展状态。"
    ))
    add_before(anchor, (
        "开展本课题的实践意义在于：以已上线运行的校园信息平台为载体，把讲座、竞赛、志愿服务、社团活动、AI 学习资源、影像资料等内容纳入统一数据入口，降低师生查找校园机会的成本，并通过智能检索与个性化推荐提高信息触达效率。"
        "项目不是另起一个展示型网站，而是在真实用户和真实数据基础上补齐数据治理、检索索引、推荐反馈、组织主体归属等关键环节。"
    ))
    add_before(anchor, (
        "开展本课题的研究意义在于：将平台运行中自然沉淀的资源覆盖度、检索行为、活动热度、报名收藏、用户反馈、组织活跃度等指标转化为可分析的数据集，探索一套轻量、可落地的校园服务质量观察方法。"
        "该方法可辅助学校发现信息供给不均衡、活动触达不足、学生需求变化等问题，为校园服务优化、学生综合发展支持和教育质量动态评估提供数据依据。"
    ))
    add_before(anchor, (
        "同时，本课题依托浙江大学多校区、多学院、多社团和校企合作资源开展试运行，具有样本丰富、迭代快、反馈真实的优势。"
        "项目形成的数据规范、平台功能、运营机制和评估报告，可为同类高校推进智慧校园建设、校园信息治理和学生发展支持提供可参考的本科生创新实践案例。"
    ))


def section_four(doc, anchor, route_img):
    add_before(anchor, "四、课题研究内容和工作方案", size=14, bold=True, font_name="黑体")
    add_before(anchor, "（一）研究内容", size=12, bold=True, font_name="黑体")
    items = [
        "1. 多源校园信息融合研究：围绕学院通知、社团活动、学术讲座、科创竞赛、志愿服务、影像资料等来源，建立统一字段和分类体系，解决来源分散、格式不一、重复发布和历史数据难沉淀的问题。",
        "2. 智能检索与精准推送平台升级：在现有平台基础上优化全站搜索、活动筛选、AI 活动推荐和用户偏好读取机制，使学生能够按时间、地点、类别、面向对象、综素/工时等条件快速发现相关机会。",
        "3. 校园服务质量指标构建：基于资源覆盖、检索命中、用户点击收藏、报名转化、活动热度、组织发布活跃度和反馈意见等数据，形成可解释、可持续更新的校园服务质量观察指标。",
        "4. 共建运营与制度化验证：完善学生、社团、学院和合作企业共同发布、审核、归属和反馈的流程，形成数据管理规范、平台运营办法和阶段性分析报告。"
    ]
    for item in items:
        add_before(anchor, item)

    add_before(anchor, "（二）技术路线", size=12, bold=True, font_name="黑体")
    add_before(anchor, (
        "本课题采用“数据来源识别—采集治理—资源索引—智能服务—运行验证”的路线推进。AI 只作为解析、检索和推荐环节的工具，不替代数据标准、平台逻辑和评估指标设计。"
    ))
    add_image_before(doc, anchor, route_img, 6.3, "图 1  本课题技术路线图")
    details = [
        "1. 数据来源与采集对象：采集对象包括校内官网通知、学院/社团公众号推文、活动发布表单、AI 社区内容、影像资料及用户行为日志。核心字段包括标题、时间、地点、主办主体、面向对象、活动类型、综素/工时信息、报名链接、浏览/搜索/收藏/反馈记录等。",
        "2. 数据清洗与标准化：采用自动化采集、人工上传和后台审核相结合的方式，使用大模型解析与规则校验提取结构化字段；对重复活动、缺失字段、过期信息和分类不一致问题进行去重、补全和标准化。平台侧以 Node.js/Express 后端、SQLite 数据库和 React 前端承载数据录入、审核、展示与检索。",
        "3. 资源索引与智能服务：建立活动库、AI 社区、影像库和组织主体档案，形成统一资源检索索引；利用关键词扩展、标签匹配、结构化筛选和轻量向量特征提升全站检索召回；活动推荐模块读取用户画像和活动特征，输出推荐理由、置信度和可追踪反馈。",
        "4. 指标计算与分析：围绕资源覆盖度、信息更新及时性、检索命中率、推荐点击/收藏/报名转化、组织发布活跃度、用户反馈满意度等指标，定期生成校园数据分析简报，为教育质量观察和校园服务优化提供依据。",
        "5. 运行验证与安全边界：依托现有线上平台进行试运行，采用日志统计、用户反馈、功能冒烟测试和阶段性人工复核验证效果；对用户数据采取最小化采集、权限隔离、匿名化统计和备份机制，避免将个人隐私直接用于公开评估。"
    ]
    for item in details:
        add_before(anchor, item)

    add_before(anchor, "（三）进度计划", size=12, bold=True, font_name="黑体")
    plan = [
        "第一阶段（2026.09—2026.11）：完成校园信息来源梳理、字段标准和分类规范设计；整理现有平台数据质量问题，形成平台升级清单和评估指标初稿。",
        "第二阶段（2026.12—2027.03）：完成采集、清洗、索引、主体归属和推荐模块迭代；接入活动、AI 社区、影像库等核心资源，形成可试运行的平台版本。",
        "第三阶段（2027.04—2027.06）：开展线上试运行与数据分析，收集用户反馈，优化检索准确率、推荐解释性和数据看板；修正指标体系和运营流程。",
        "第四阶段（2027.07—2027.08）：整理平台成果、研究报告、数据规范、运营制度和案例材料，完成结题验收准备。"
    ]
    for item in plan:
        add_before(anchor, item)


def section_five(doc, anchor, collage_img):
    add_before(anchor, "五、基础条件和优势", size=14, bold=True, font_name="黑体")
    add_before(anchor, (
        "项目已具备真实平台、真实用户和真实校园数据基础。现有“拓途浙享”平台已完成活动聚合、AI 社区、影像库、上传审核、全站搜索等基础模块，累计注册用户超过 2500 人，日访问量稳定破千，收录校园活动 700 余项，并在首届“浙客松”AI 黑客松等实践中形成了校内外合作基础。"
    ))
    add_table_before(doc, anchor, [
        ["现有模块", "已有基础", "本课题重点提升方向"],
        ["活动聚合", "活动分类筛选、时间地点展示、创建活动、日历/分享等基础能力已上线。", "提升多源采集稳定性、字段完整率、活动去重和过期信息治理。"],
        ["AI 社区", "已包含技术分享、求助、社群、新闻等内容入口。", "强化内容标签、学习资源归类和活动/项目之间的关联推荐。"],
        ["影像库", "已支持现场照片、视频记录和分类浏览。", "补齐影像元数据、活动关联和成果佐证归档。"],
        ["主体归属", "已具备个人、社团、学院、企业等 profile 主体和发布/主办归属雏形。", "完善组织别名识别、历史数据迁移、审核日志和共建运营规则。"],
        ["智能检索推荐", "已有全站搜索、活动推荐助手和资源索引基础。", "提升索引覆盖、推荐解释性、反馈闭环和评估指标输出。"],
    ], widths=[1.15, 2.35, 2.55])
    add_caption_before(anchor, "表 1  现有平台基础与本课题提升方向")
    add_image_before(doc, anchor, collage_img, 6.3, "图 2  现有平台主要功能界面")
    add_before(anchor, (
        "需要说明的是，现有平台只是课题开展的基础条件，并不等同于本课题已经完成。当前仍存在多源信息采集质量不稳定、不同来源字段标准不一致、组织主体归属不够清晰、智能检索和推荐缺少持续评估、平台运行数据尚未系统转化为教育质量观察指标等问题。"
        "本次申报经费主要用于平台升级、数据治理、模型接口、服务器扩容、用户调研和成果编制，目标是把已有应用基础提升为可验证、可总结、可推广的研究成果。"
    ))
    add_before(anchor, (
        "团队方面，课题组成员已长期参与平台研发、数据维护、活动运营和校企合作对接，具备全栈开发、数据分析、新媒体运营和项目组织经验。学校完善的网络环境、服务器资源和浙江大学丰富的校园活动场景，也为课题持续试运行和迭代验证提供了保障。"
    ))


def section_six(doc, anchor):
    add_before(anchor, "六、预期成果和提交方式", size=14, bold=True, font_name="黑体")
    add_before(anchor, "（一）预期成果", size=12, bold=True, font_name="黑体")
    outcomes = [
        "1. 平台成果：完成“拓途浙享”校园信息智能检索平台升级版本，重点实现多源活动数据采集与标准化、全站资源检索、活动个性化推荐、组织主体归属、影像与成果资料归档、基础数据统计等功能，并保障电脑端和移动端稳定访问。",
        "2. 数据与研究成果：形成《基于多源数据融合的高校教育质量评估体系研究报告》1 份，围绕资源覆盖、信息触达、活动参与、组织活跃、用户反馈等指标开展实证分析；形成若干校园数据分析简报，展示平台运行数据对校园服务优化的支撑作用。",
        "3. 制度与文档成果：形成《校园信息数据管理规范》《平台运营管理办法》《用户信息安全管理制度》以及接口说明、数据库设计、系统运维和故障排查文档，支撑平台后续常态化运营。",
        "4. 实践应用成果：在浙江大学校园场景中完成试运行，沉淀用户反馈、平台截图、活动案例、校企交流材料和阶段性报告，形成面向同类高校的实施清单与案例材料。"
    ]
    for item in outcomes:
        add_before(anchor, item)

    add_before(anchor, "（二）提交方式", size=12, bold=True, font_name="黑体")
    submissions = [
        "系统平台：提交可正常访问的平台地址、核心功能说明、运行截图和后台数据统计结果，接受线上核验。",
        "文字资料：提交研究报告、数据规范、运营制度、技术文档、阶段性报告和案例材料，提供电子版和按要求装订的纸质版。",
        "佐证材料：提交平台运行数据报表、用户反馈记录、活动合作记录、赛事/社区影像资料和校企交流材料。"
    ]
    for item in submissions:
        add_before(anchor, item)

    add_before(anchor, "（三）创新点", size=12, bold=True, font_name="黑体")
    add_before(anchor, (
        "1. 以日常校园服务数据支撑教育质量动态观察。课题不依赖一次性问卷或人工汇总，而是从学生真实检索、浏览、收藏、报名、反馈和组织发布行为中提取指标，使校园服务质量和学生发展机会的观察更及时、更贴近实际场景。"
    ))
    add_before(anchor, (
        "2. 将信息治理、智能服务和组织共建放在同一平台闭环验证。项目把多源数据采集、字段标准化、资源索引、AI 检索推荐、发布主体归属和运营制度放入同一个线上平台中试运行，能够同时验证技术可行性、用户体验和管理制度，避免只做算法演示或只做信息展示。"
    ))


def main():
    route_img = ASSET_DIR / "technical_route.png"
    collage_img = ASSET_DIR / "platform_collage.png"
    make_route_chart(route_img)
    make_platform_collage(collage_img)

    doc = Document(BASE)
    update_summary_table(doc)
    replace_section(doc, "三、课题研究的目的和意义", "四、课题研究内容和工作方案", lambda anchor: section_three(doc, anchor))
    replace_section(doc, "四、课题研究内容和工作方案", "五、基础条件和优势", lambda anchor: section_four(doc, anchor, route_img))
    replace_section(doc, "五、基础条件和优势", "六、预期成果和提交方式", lambda anchor: section_five(doc, anchor, collage_img))
    replace_section(doc, "六、预期成果和提交方式", "七、课题经费使用计划", lambda anchor: section_six(doc, anchor))

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
